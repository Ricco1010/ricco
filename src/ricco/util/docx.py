import os

import pandas as pd

from .decorator import print_doc
from .dt import DT


class Docx:
  """Docx"""

  def set_default_style(self):
    """格式初始化"""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt
    from docx.shared import RGBColor

    self.doc = Document()
    # 设置默认字体
    self.doc.styles['Normal'].font.name = u'微软雅黑'
    self.doc.styles['Normal']._element.rPr.rFonts.set(
        qn('w:eastAsia'), u'微软雅黑')
    self.doc.styles['Normal'].font.size = Pt(10.5)
    self.doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
    self.create_time = DT().today

  def add_table_from_df(self, df: pd.DataFrame, newline=False):
    """将dataframe作为表格写入"""
    t = self.doc.add_table(df.shape[0] + 1, df.shape[1], style='Table Grid')
    for j in range(df.shape[-1]):
      t.cell(0, j).text = df.columns[j]
    for i in range(df.shape[0]):
      for j in range(df.shape[-1]):
        t.cell(i + 1, j).text = str(df.values[i, j])
    if newline:
      self.add_normal_p('')

  def add_title(self, text: str, n: int):
    """添加标题"""
    self.doc.add_heading(text, level=n)

  def add_title_0(self, text: str):
    """添加大标题"""
    self.add_title(text, 0)

  def add_title_1(self, text: str):
    """添加一级标题"""
    self.add_title(text, 1)

  def add_title_2(self, text: str):
    """添加二级标题"""
    self.add_title(text, 2)

  def add_title_3(self, text: str):
    """添加三级标题"""
    self.add_title(text, 3)

  def add_title_4(self, text: str):
    """添加四级标题"""
    self.add_title(text, 4)

  def add_normal_p(self, text: str):
    """添加普通段落"""
    self.doc.add_paragraph(text)

  def add_paragraph_color(self, text: str, rgb: (list, tuple)):
    """添加自定义颜色的段落"""
    from docx.shared import RGBColor

    p = self.doc.add_paragraph('')
    p.add_run(text).font.color.rgb = RGBColor(*rgb)

  def add_paragraph_red(self, text: str):
    """添加红色字体的段落"""
    from docx.shared import RGBColor

    p = self.doc.add_paragraph('')
    p.add_run(text).font.color.rgb = RGBColor(250, 0, 0)

  def add_bullet_list(self, text: str):
    """添加无序列表"""
    self.doc.add_paragraph(text, style='List Bullet')

  def add_order_list(self, text: str):
    """添加有序列表"""
    self.doc.add_paragraph(text, style='List Number')

  def add_intense_quote(self, text: str):
    """添加高亮段落"""
    self.doc.add_paragraph(text, style='Intense Quote')

  @print_doc
  def save(self, file_path: str):
    """保存文件至word文档"""
    self.doc.save(file_path)
    print('文件保存至：', os.path.abspath(file_path))

  def add_hist_from_data(self, data: list, newline=False):
    """通过数据生成并添加直方图"""
    from docx.shared import Inches
    from matplotlib import pyplot as plt
    plt.figure(figsize=(12, 4))
    plt.style.use('seaborn')
    plt.hist(data)
    plt.savefig('image.png')
    self.doc.add_picture('image.png', width=Inches(6))
    if os.path.exists('image.png'):
      os.remove('image.png')
    if newline:
      self.add_normal_p('')
