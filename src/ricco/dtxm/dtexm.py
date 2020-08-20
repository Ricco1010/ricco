# encoding:GBK
from docx import Document
from docx.shared import Inches
from ricco import rdf
import pandas as pd
import numpy as np
import os
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from ricco.util import fn
from ricco.dtxm.wiki import class_dic
from ricco.util import col_round
from ricco.util import ensure_list
import matplotlib.pyplot as plt


class Dtexm(object):
    def __init__(self, filename, cols_list: list = None):
        self.filename = filename
        self.doc = Document()
        # 设置默认字体
        self.doc.styles['Normal'].font.name = u'微软雅黑'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
        self.doc.styles['Normal'].font.size = Pt(10.5)
        self.doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
        self.doc.add_heading('数据检测报告', 0)

        if isinstance(self.filename, str):
            self.add_intense_quote(f'Document：{self.filename}')
            self.df = rdf(self.filename)
        elif isinstance(self.filename, pd.DataFrame):
            self.df = self.filename
        else:
            raise ValueError('请输入Dataframe或路径')
        if cols_list != None:
            cols_list = ensure_list(cols_list)
            self.df = self.df[[cols_list]]
        self.lenth = len(self.df)

    def add_df2table(self, table_df):
        '''将dataframe作为表格写入'''
        t = self.doc.add_table(table_df.shape[0] + 1, table_df.shape[1], style='Table Grid')
        for j in range(table_df.shape[-1]):
            t.cell(0, j).text = table_df.columns[j]
        for i in range(table_df.shape[0]):
            for j in range(table_df.shape[-1]):
                t.cell(i + 1, j).text = str(table_df.values[i, j])

    def add_normal_p(self, text):
        '''添加普通段落'''
        self.doc.add_paragraph(text)

    def add_bullet_list(self, text):
        '''添加无序列表'''
        self.doc.add_paragraph(text, style='List Bullet')

    def add_order_list(self, text):
        '''添加有序列表'''
        self.doc.add_paragraph(text, style='List Number')

    def add_intense_quote(self, text):
        self.doc.add_paragraph(text, style='Intense Quote')

    def add_title(self, text, n):
        '''添加标题'''
        self.doc.add_heading(text, level=n)

    def basic(self):
        '''数据基础信息描述'''
        # 列名
        self.add_bullet_list('列名：')
        self.add_normal_p('，'.join(self.df.columns))
        # 文件size
        self.add_bullet_list('文件尺寸：')
        self.add_normal_p(f'行：{self.df.shape[0]}，列{self.df.shape[1]}')
        # 列类型
        self.add_bullet_list('数据类型：')
        self.col_types = pd.DataFrame(self.df.dtypes, columns=['类型']).reset_index().rename(columns={'index': '列名'})
        self.add_df2table(self.col_types)

    def serise_describe(self, col):
        '''数值型列的描述性统计'''
        desc = pd.DataFrame(self.df[col].describe().reset_index())
        desc = desc.rename(columns={'index': '分类', col: '值'})
        desc['分类'] = desc['分类'].replace(to_replace=class_dic)
        skew_add = pd.DataFrame({'分类': '偏度系数',
                                 '值': [self.df[col].skew()]})
        kurt_add = pd.DataFrame({'分类': '峰度系数',
                                 '值': [self.df[col].kurt() - 3]})

        null_num = self.lenth - desc.loc[desc['分类'] == '计数', '值'][0]
        null_rate = null_num / self.lenth

        null_add = pd.DataFrame({'分类': '缺失数',
                                 '值': [null_num]})
        null_rate_add = pd.DataFrame({'分类': '缺失率',
                                      '值': [null_rate]})
        desc = desc.append(skew_add, sort=False)
        desc = desc.append(kurt_add, sort=False)
        desc = desc.append(null_add, sort=False)
        desc = desc.append(null_rate_add, sort=False)
        return desc

    def object_describe(self, col):
        desc = pd.DataFrame(self.df[col].value_counts().reset_index())
        if len(desc) > 20:
            desc = desc.rename(columns={'index': '分类_Top15', col: '数量'})
            desc['分类_Top15'] = desc['分类_Top15'].replace(to_replace=class_dic)
            return desc.head(15)
        else:
            desc = desc.rename(columns={'index': '分类', col: '数量'})
            desc['分类'] = desc['分类'].replace(to_replace=class_dic)
            return desc

    def is_float(self, col):
        def try2float(x):
            try:
                return float(x)
            except ValueError:
                return None

        length = len(self.df[~self.df[col].isna()])
        null_num = len(self.df[~self.df[col].apply(lambda x: try2float(x)).isna()])
        rates = null_num / length
        if rates >= 0.8:
            text = f' {col} 列有超过{int(rates * 100)}%的值为数值型的数据'
            p = self.doc.add_paragraph('')
            p.add_run(text).font.color.rgb = RGBColor(250, 0, 0)

    def is_date(self, col):
        length = len(self.df[~self.df[col].isna()])
        null_num = len(self.df[~pd.to_datetime(self.df[col], errors='coerce').isna()])
        rates = null_num / length
        if rates >= 0.8:
            text = f' {col} 列有超过{int(rates * 100)}%的值为日期格式的数据'
            p = self.doc.add_paragraph('')
            p.add_run(text).font.color.rgb = RGBColor(250, 0, 0)

    def hist_plot(self, col):
        plt.figure(figsize=(12, 4))
        plt.style.use('seaborn')
        data = self.df[~self.df[col].isna()][col].values
        plt.hist(data)
        plt.savefig('image.png')
        self.doc.add_picture('image.png', width=Inches(6))
        if os.path.exists('image.png'):
            os.remove('image.png')

    def col_by_col(self):
        '''逐列检测'''
        for col in self.df.columns:
            self.add_bullet_list(col)
            if (self.df[col].dtype == 'int64') | (self.df[col].dtype == 'float64'):
                desc_df = self.serise_describe(col)
                desc_df = col_round(desc_df, '值')
                self.add_df2table(desc_df)
                self.hist_plot(col)
                self.add_normal_p('')
            else:
                desc_df = self.object_describe(col)
                self.add_df2table(desc_df)
                self.add_normal_p('')
                self.is_float(col)
                self.is_date(col)

    def save(self, savefilename: str = None):
        '''保存文件至word文档'''
        if savefilename == None:
            if isinstance(self.filename, pd.DataFrame):
                raise FileNotFoundError('请输入要保存的文件路径')
            savefilename = fn(self.filename) + '-' '检测报告.docx'
        self.doc.save(savefilename)
        print('文件保存至', os.path.abspath(savefilename))

    def examine_all(self):
        '''整套流程'''
        self.basic()
        self.col_by_col()
        self.save()


if __name__ == '__main__':
    doc = Dtexm('sample.csv')
    doc.examine_all()
