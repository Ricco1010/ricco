# encoding:GBK
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import RGBColor


class _Docx(object):

    def __init__(self):
        self.doc = Document()
        # 设置默认字体
        self.doc.styles['Normal'].font.name = u'微软雅黑'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
        self.doc.styles['Normal'].font.size = Pt(10.5)
        self.doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
        self.doc.add_heading('数据检测报告', 0)

    def add_df2table(self, table_df):
        '''将dataframe作为表格写入'''
        t = self.doc.add_table(table_df.shape[0] + 1, table_df.shape[1], style='Table Grid')
        for j in range(table_df.shape[-1]):
            t.cell(0, j).text = table_df.columns[j]
        for i in range(table_df.shape[0]):
            for j in range(table_df.shape[-1]):
                t.cell(i + 1, j).text = str(table_df.values[i, j])

    def add_normal_p(self, text):
        '''添加普通段落'''
        self.doc.add_paragraph(text)

    def add_bullet_list(self, text):
        '''添加无序列表'''
        self.doc.add_paragraph(text, style='List Bullet')

    def add_order_list(self, text):
        '''添加有序列表'''
        self.doc.add_paragraph(text, style='List Number')

    def add_intense_quote(self, text):
        self.doc.add_paragraph(text, style='Intense Quote')

    def add_title(self, text, n):
        '''添加标题'''
        self.doc.add_heading(text, level=n)
